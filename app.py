import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
import time
import random

# ------------------------------------------------------------
# ГЛОБАЛЬНОЕ СОСТОЯНИЕ ДЛЯ ШТРАФНОЙ ФУНКЦИИ
# ------------------------------------------------------------
PENALTY_STATE = {}

# ------------------------------------------------------------
# 1. ВЫЧИСЛИТЕЛЬНОЕ ЯДРО (ОБНОВЛЁННОЕ)
# ------------------------------------------------------------
def calculate_piket(H, B, c, phi, gamma, f, UGW):
    """
    Расчёт параметров для одного пикета.
    Вертикальное давление: полное горное давление sigma_v = gamma_kN * H.
    Боковое давление: по методу Фотта с использованием нового sigma_v.
    Пластическая зона: по эффективному напряжению sigma_0_eff = gamma_kN * H - u.
    """
    # ---------- Защита от экстремальных значений ----------
    if c < 1.0:
        c = 1.0
    if phi < 5.0:
        phi = 5.0
    if f < 0.5:
        f = 0.5

    gamma_kN = gamma * 9.81                     # кН/м³
    phi_rad = np.radians(phi)
    sin_phi = np.sin(phi_rad)
    cos_phi = np.cos(phi_rad)
    tan_phi = np.tan(phi_rad)
    if tan_phi < 0.1:
        tan_phi = 0.1

    # ---------- Вертикальное давление (полное горное) ----------
    sigma_v = gamma_kN * H

    # ---------- (Информационно) высота свода обрушения по Протодьяконову ----------
    tg_45_minus = np.tan(np.pi/4 - phi_rad/2)
    h_arch = (B + 2 * H * tg_45_minus) / (2 * f) if f > 0 else 0.0
    if H > 80:
        h_arch = min(h_arch, H / 3.0)

    # ---------- Боковое давление (метод Фотта) ----------
    K_a = np.tan(np.pi/4 - phi_rad/2)**2
    sqrt_K_a = np.sqrt(K_a)
    sigma_h = sigma_v * K_a - 2 * c * sqrt_K_a
    if sigma_h < 0:
        sigma_h = 0.0

    # ---------- Поровое давление ----------
    gamma_w = 9.81
    if UGW < H:
        u = gamma_w * (H - UGW)
    else:
        u = 0.0

    # ---------- Пластическая зона (по эффективным напряжениям) ----------
    R_0 = B / 2.0
    sigma_0_eff = gamma_kN * H - u
    if sigma_0_eff < 0:
        sigma_0_eff = 0.0

    p_i = 0.0

    if sin_phi == 0:
        if sigma_0_eff > c:
            plastic_zone = True
            plastic_radius = R_0 * (sigma_0_eff / c) if c > 0 else R_0 * 10.0
        else:
            plastic_zone = False
            plastic_radius = 0.0
    else:
        cot_phi = 1.0 / tan_phi
        term1 = (sigma_0_eff + c * cot_phi) * (1 - sin_phi)
        term2 = p_i + c * cot_phi

        if term2 <= 0:
            plastic_zone = True
            plastic_radius = R_0 * 100.0
        else:
            ratio = term1 / term2
            exponent = (1 - sin_phi) / (2 * sin_phi)
            if ratio > 1:
                R_p = R_0 * (ratio ** exponent)
                plastic_zone = True
                plastic_radius = R_p
            else:
                plastic_zone = False
                plastic_radius = 0.0

    return {
        'sigma_v': sigma_v,
        'sigma_h': sigma_h,
        'u': u,
        'plastic_zone': plastic_zone,
        'plastic_radius': plastic_radius,
        'h_arch': h_arch
    }


def apply_penalty(results, excess_threshold=0.9, stability_threshold=0.8, include_plastic=False):
    """
    Асимметричная штрафная функция с настраиваемыми порогами.
    - excess_threshold: порог для u / sigma_v (по умолчанию 0.9)
    - stability_threshold: порог для stability_ratio (по умолчанию 0.8)
    - include_plastic: если True, то пластика участвует в критерии красного
    """
    piket = results.get('piket')
    if piket is None:
        raise ValueError("Missing 'piket' in results")

    if piket not in PENALTY_STATE:
        PENALTY_STATE[piket] = {'is_red': False}
    is_red_prev = PENALTY_STATE[piket]['is_red']

    sigma_v = results.get('sigma_v', 0)
    u = results.get('u', 0)
    plastic = results.get('plastic_zone', False)

    # 1. Превышение порового давления с настраиваемым порогом
    exceeds_u = u > excess_threshold * sigma_v

    # 2. Коэффициент устойчивости (по эффективным напряжениям)
    stability_ratio = 1.0
    if 'c' in results and 'phi' in results and sigma_v > 0:
        phi_rad = np.radians(results['phi'])
        sigma_eff = sigma_v - u
        if sigma_eff < 0:
            sigma_eff = 0
        tau_lim = results['c'] + sigma_eff * np.tan(phi_rad)
        sigma_h = results.get('sigma_h', 0)
        tau_act = (sigma_v - sigma_h) / 2 if sigma_v > sigma_h else 0.5
        if tau_act > 0:
            stability_ratio = tau_lim / tau_act
        else:
            stability_ratio = float('inf')
    low_stability = stability_ratio < stability_threshold

    # 3. Логика красного с учётом include_plastic
    if include_plastic:
        is_red_now = (exceeds_u and low_stability) or (plastic and exceeds_u)
    else:
        is_red_now = exceeds_u and low_stability

    # Асимметричная память
    if is_red_prev:
        is_red_final = True
        color = 'red'
    else:
        if is_red_now:
            is_red_final = True
            color = 'red'
        else:
            is_red_final = False
            # Жёлтый: любое одно из условий (exceeds_u, low_stability, plastic)
            if exceeds_u or low_stability or plastic:
                color = 'yellow'
            else:
                color = 'green'

    PENALTY_STATE[piket]['is_red'] = is_red_final

    # Риск
    if color == 'red':
        risk_score = 1.0
    elif color == 'yellow':
        risk_score = 0.5
    else:
        risk_score = 0.0

    if not np.isfinite(risk_score):
        risk_score = 1e9

    results['color'] = color
    results['risk_score'] = risk_score
    results['is_red'] = is_red_final
    results['stability_ratio'] = stability_ratio
    results['flag_exceeds_u'] = exceeds_u
    results['flag_low_stability'] = low_stability
    results['flag_plastic'] = plastic

    return results


def detect_anomalies(df, window_size=5, sigma_threshold=3.0):
    df = df.copy()
    df = df.sort_values('piket').reset_index(drop=True)
    n = len(df)

    for col in df.columns:
        if df[col].isnull().any():
            col_mean = df[col].mean(skipna=True)
            if np.isnan(col_mean):
                col_mean = 0.0
            df[col].fillna(col_mean, inplace=True)

    df['anomaly_geomech'] = False
    df['anomaly_hydro'] = False
    df['anomaly_geol'] = False

    effective_window = min(window_size, n)
    if effective_window < 3:
        df['overall_risk'] = df.get('risk_score', 0)
        return df

    for i in range(n):
        start = max(0, i - effective_window // 2)
        end = min(n, i + effective_window // 2 + 1)
        window_indices = list(range(start, end))
        if len(window_indices) < 3:
            continue

        sigma_v_win = df.loc[window_indices, 'sigma_v'].values
        sigma_h_win = df.loc[window_indices, 'sigma_h'].values
        plastic_win = df.loc[window_indices, 'plastic_zone'].astype(int).values

        sigma_v_i = df.loc[i, 'sigma_v']
        sigma_h_i = df.loc[i, 'sigma_h']
        plastic_i = df.loc[i, 'plastic_zone']

        mean_v = np.mean(sigma_v_win)
        std_v = np.std(sigma_v_win)
        if std_v > 0 and abs(sigma_v_i - mean_v) > sigma_threshold * std_v:
            df.loc[i, 'anomaly_geomech'] = True

        mean_h = np.mean(sigma_h_win)
        std_h = np.std(sigma_h_win)
        if std_h > 0 and abs(sigma_h_i - mean_h) > sigma_threshold * std_h:
            df.loc[i, 'anomaly_geomech'] = True

        plastic_fraction = np.mean(plastic_win)
        if (plastic_i == 1 and plastic_fraction < 0.3) or (plastic_i == 0 and plastic_fraction > 0.7):
            df.loc[i, 'anomaly_geomech'] = True

    if 'H' not in df.columns:
        df['H'] = 0
    for i in range(n):
        if df.loc[i, 'UGW'] > 0.7 * df.loc[i, 'H']:
            df.loc[i, 'anomaly_hydro'] = True

        start = max(0, i - effective_window // 2)
        end = min(n, i + effective_window // 2 + 1)
        window_indices = list(range(start, end))
        if len(window_indices) < 3:
            continue
        u_win = df.loc[window_indices, 'u'].values
        u_i = df.loc[i, 'u']
        mean_u = np.mean(u_win)
        std_u = np.std(u_win)
        if std_u > 0 and abs(u_i - mean_u) > sigma_threshold * std_u:
            df.loc[i, 'anomaly_hydro'] = True

    for i in range(n):
        f_i = df.loc[i, 'f']
        phi_i = df.loc[i, 'phi']
        c_i = df.loc[i, 'c']

        if f_i < 3.0 and phi_i > 35.0 and c_i < 20.0:
            df.loc[i, 'anomaly_geol'] = True

        start = max(0, i - effective_window // 2)
        end = min(n, i + effective_window // 2 + 1)
        window_indices = list(range(start, end))
        if len(window_indices) < 3:
            continue
        f_win = df.loc[window_indices, 'f'].values
        phi_win = df.loc[window_indices, 'phi'].values
        c_win = df.loc[window_indices, 'c'].values

        mean_f = np.mean(f_win)
        std_f = np.std(f_win)
        mean_phi = np.mean(phi_win)
        std_phi = np.std(phi_win)
        mean_c = np.mean(c_win)
        std_c = np.std(c_win)

        if (std_f > 0 and abs(f_i - mean_f) > sigma_threshold * std_f) or \
           (std_phi > 0 and abs(phi_i - mean_phi) > sigma_threshold * std_phi) or \
           (std_c > 0 and abs(c_i - mean_c) > sigma_threshold * std_c):
            df.loc[i, 'anomaly_geol'] = True

    if 'risk_score' in df.columns:
        df['overall_risk'] = df['risk_score']
    else:
        df['overall_risk'] = (df['anomaly_geomech'].astype(int) +
                              df['anomaly_hydro'].astype(int) +
                              df['anomaly_geol'].astype(int) +
                              df['plastic_zone'].astype(int))

    return df


def compute_full_table(df_input, excess_threshold=0.9, stability_threshold=0.8, include_plastic=False):
    PENALTY_STATE.clear()
    results_list = []
    for idx, row in df_input.iterrows():
        piket = row['piket']
        H = row['H']
        B = row['B']
        c = row['c']
        phi = row['phi']
        gamma = row['gamma']
        f = row['f']
        UGW = row['UGW']
        res = calculate_piket(H, B, c, phi, gamma, f, UGW)
        res['piket'] = piket
        res['H'] = H
        res['B'] = B
        res['c'] = c
        res['phi'] = phi
        res['gamma'] = gamma
        res['f'] = f
        res['UGW'] = UGW
        res = apply_penalty(res, excess_threshold, stability_threshold, include_plastic)
        results_list.append(res)

    df_results = pd.DataFrame(results_list)
    df_full = detect_anomalies(df_results)
    return df_full


# ------------------------------------------------------------
# 2. ВИЗУАЛИЗАЦИЯ (plotly)
# ------------------------------------------------------------
def plot_longitudinal_profile(df):
    fig = go.Figure()
    surface = df['H'] + np.random.normal(0, 2, len(df))
    fig.add_trace(go.Scatter(
        x=df['piket'],
        y=surface,
        mode='lines',
        name='Рельеф (поверхность)',
        line=dict(color='brown', width=2),
        fill='tozeroy',
        fillcolor='rgba(139, 69, 19, 0.2)'
    ))
    fig.add_trace(go.Scatter(
        x=df['piket'],
        y=df['UGW'],
        mode='lines+markers',
        name='УГВ',
        line=dict(color='blue', dash='dash'),
        marker=dict(size=6)
    ))
    tunnel_depth = surface - df['H']
    fig.add_trace(go.Scatter(
        x=df['piket'],
        y=tunnel_depth,
        mode='lines+markers',
        name='Тоннель (ось)',
        line=dict(color='red', width=3),
        marker=dict(size=8)
    ))
    fig.update_layout(
        title='Продольный разрез трассы',
        xaxis_title='Пикетаж (м)',
        yaxis_title='Отметка (м)',
        hovermode='x unified',
        template='plotly_white',
        height=600
    )
    return fig


def plot_picket_epures(df, piket_index):
    row = df[df['piket'] == piket_index]
    if row.empty:
        return go.Figure()
    row = row.iloc[0]
    depth = np.linspace(0, row['H'], 50)
    gamma_kN = row['gamma'] * 9.81
    sigma_v_vals = gamma_kN * depth
    K0 = 1 - np.sin(np.radians(row['phi']))
    sigma_h_vals = K0 * sigma_v_vals
    u_vals = np.zeros_like(depth)
    if row['UGW'] < row['H']:
        mask = depth > row['UGW']
        u_vals[mask] = 9.81 * (depth[mask] - row['UGW'])

    fig = make_subplots(
        rows=1, cols=4,
        subplot_titles=('σ_v (кПа)', 'σ_h (кПа)', 'Поровое u (кПа)', 'Зона пластики')
    )
    fig.add_trace(go.Scatter(x=sigma_v_vals, y=depth, mode='lines', name='σ_v', line=dict(color='blue')), row=1, col=1)
    fig.add_trace(go.Scatter(x=sigma_h_vals, y=depth, mode='lines', name='σ_h', line=dict(color='green')), row=1, col=2)
    fig.add_trace(go.Scatter(x=u_vals, y=depth, mode='lines', name='u', line=dict(color='red')), row=1, col=3)
    plastic_val = 1 if row['plastic_zone'] else 0
    fig.add_trace(go.Bar(x=[plastic_val], y=[row['H']], name='Пластика', orientation='v',
                         marker_color='orange' if plastic_val else 'lightgreen'),
                  row=1, col=4)

    for i in range(1, 4):
        fig.update_yaxes(title_text='Глубина (м)', autorange='reversed', row=1, col=i)
    fig.update_xaxes(title_text='Напряжение (кПа)', row=1, col=1)
    fig.update_xaxes(title_text='Давление (кПа)', row=1, col=2)
    fig.update_xaxes(title_text='Давление (кПа)', row=1, col=3)
    fig.update_xaxes(title_text='Пластика', row=1, col=4)

    fig.update_layout(
        height=550,
        showlegend=False,
        template='plotly_white'
    )
    return fig


def plot_ccm_curves(df, selected_pikets=None):
    if selected_pikets is None:
        selected_pikets = df['piket'].iloc[:3].tolist()
    fig = go.Figure()
    for p in selected_pikets:
        row = df[df['piket'] == p]
        if row.empty:
            continue
        row = row.iloc[0]
        eps = np.linspace(0, 0.1, 100)
        sigma_max = row['sigma_v'] * 1.2
        load = sigma_max * (1 - np.exp(-eps * 30))
        load += np.random.normal(0, 0.02*sigma_max, size=len(eps))
        fig.add_trace(go.Scatter(x=eps, y=load, mode='lines', name=f'Пикет {p}', line=dict(width=2)))
    fig.update_layout(
        title='Кривые взаимодействия (CCM)',
        xaxis_title='Деформация ε',
        yaxis_title='Нагрузка (кПа)',
        hovermode='x',
        template='plotly_white',
        height=450
    )
    return fig


def plot_correlation_matrix(df):
    """
    Интерактивная корреляционная матрица для числовых параметров.
    Отображает корреляции между входными данными и результатами расчёта.
    """
    df_corr = df.copy()
    df_corr['plastic_int'] = df_corr['plastic_zone'].astype(int)

    numeric_cols = ['H', 'B', 'c', 'phi', 'gamma', 'f', 'UGW', 
                    'sigma_v', 'sigma_h', 'u', 'risk_score', 'plastic_int']
    available_cols = [col for col in numeric_cols if col in df_corr.columns]
    corr_matrix = df_corr[available_cols].corr()

    rename_map = {
        'H': 'Глубина H',
        'B': 'Диаметр B',
        'c': 'Сцепление c',
        'phi': 'Угол трения φ',
        'gamma': 'Плотность γ',
        'f': 'Крепость f',
        'UGW': 'УГВ',
        'sigma_v': 'σ_v',
        'sigma_h': 'σ_h',
        'u': 'Поровое u',
        'risk_score': 'Риск',
        'plastic_int': 'Пластика'
    }
    final_rename = {k: v for k, v in rename_map.items() if k in corr_matrix.columns}
    corr_matrix = corr_matrix.rename(columns=final_rename, index=final_rename)

    fig = go.Figure(data=go.Heatmap(
        z=corr_matrix.values,
        x=corr_matrix.columns,
        y=corr_matrix.index,
        colorscale='RdBu_r',
        zmin=-1,
        zmax=1,
        text=corr_matrix.round(2).values,
        texttemplate='%{text}',
        textfont={"size": 10},
        hoverongaps=False,
        colorbar=dict(title="Корреляция")
    ))
    fig.update_layout(
        title='Корреляционная матрица параметров',
        xaxis_title='Параметры',
        yaxis_title='Параметры',
        height=650,
        width=700,
        template='plotly_white'
    )
    return fig


def plot_anomaly_table(df):
    df_display = df.copy()
    def anomaly_type(row):
        types = []
        if row['anomaly_geomech']:
            types.append('Геомех.')
        if row['anomaly_hydro']:
            types.append('Гидро')
        if row['anomaly_geol']:
            types.append('Геол.')
        return ', '.join(types) if types else 'Нет'
    df_display['Тип аномалии'] = df_display.apply(anomaly_type, axis=1)
    cols = ['piket', 'H', 'sigma_v', 'sigma_h', 'u', 'plastic_zone', 'color', 'Тип аномалии']
    table_data = df_display[cols].copy()
    table_data['plastic_zone'] = table_data['plastic_zone'].map({True: 'Да', False: 'Нет'})
    fig = go.Figure(data=[go.Table(
        header=dict(
            values=list(table_data.columns),
            fill_color='paleturquoise',
            align='center',
            font=dict(size=12)
        ),
        cells=dict(
            values=[table_data[col] for col in table_data.columns],
            fill_color=[
                ['lightgreen' if c == 'green' else 'lightyellow' if c == 'yellow' else 'lightcoral'
                 for c in table_data['color']]
            ],
            align='center',
            font=dict(size=11)
        )
    )])
    fig.update_layout(
        title='Таблица аномалий',
        height=300 + 30 * len(df),
        template='plotly_white'
    )
    return fig


# ------------------------------------------------------------
# 3. ГЕНЕРАЦИЯ ОТЧЁТОВ
# ------------------------------------------------------------
def generate_docx_bytes(df, project_name="Тоннель"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph('\n' * 6)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ГЕОМЕХАНИЧЕСКИЙ СКРИНИНГ ТОННЕЛЯ')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    doc.add_paragraph('\n')
    proj = doc.add_paragraph()
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj.add_run(f'Проект: {project_name}').bold = True
    doc.add_paragraph('\n' * 4)
    date_ver = doc.add_paragraph()
    date_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_ver.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run('Версия: 1.0')
    doc.add_paragraph('\n' * 6)
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc.add_run('Дисклеймер: Финальное проектное решение принимает главный инженер проекта (ГИП). '
                 'Представленные расчёты носят рекомендательный характер и не отменяют экспертизу.')
    doc.add_page_break()

    doc.add_heading('1. Вводные параметры', level=1)
    input_cols = ['piket', 'H', 'B', 'c', 'phi', 'gamma', 'f', 'UGW']
    header_map = {
        'piket': 'Пикетаж (м)',
        'H': 'Глубина H (м)',
        'B': 'Диаметр B (м)',
        'c': 'Сцепление c (кПа)',
        'phi': 'Угол трения φ (°)',
        'gamma': 'Плотность γ (т/м³)',
        'f': 'Коэф. крепости f',
        'UGW': 'УГВ (м)'
    }
    df_input = df[input_cols].copy().rename(columns=header_map)
    table = doc.add_table(rows=1, cols=len(df_input.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df_input.columns):
        cell = table.rows[0].cells[i]
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df_input.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df_input.columns):
            row_cells[i].text = str(row[col])
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('2. Результаты расчётов', level=1)
    result_cols = ['piket', 'sigma_v', 'sigma_h', 'u', 'plastic_zone', 'color', 'risk_score']
    header_map_res = {
        'piket': 'Пикетаж (м)',
        'sigma_v': 'σ_v (кПа)',
        'sigma_h': 'σ_h (кПа)',
        'u': 'Поровое u (кПа)',
        'plastic_zone': 'Пластика',
        'color': 'Цвет риска',
        'risk_score': 'Риск'
    }
    df_res = df[result_cols].copy()
    df_res['plastic_zone'] = df_res['plastic_zone'].map({True: 'Да', False: 'Нет'})
    df_res.rename(columns=header_map_res, inplace=True)
    table_res = doc.add_table(rows=1, cols=len(df_res.columns))
    table_res.style = 'Table Grid'
    table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df_res.columns):
        cell = table_res.rows[0].cells[i]
        cell.text = col
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df_res.iterrows():
        row_cells = table_res.add_row().cells
        for i, col in enumerate(df_res.columns):
            row_cells[i].text = str(row[col])
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col == 'Цвет риска':
                color_val = str(row[col]).lower()
                if color_val == 'red':
                    row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                elif color_val == 'yellow':
                    row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 0)
                elif color_val == 'green':
                    row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
    doc.add_paragraph()

    doc.add_heading('3. Заключение', level=1)
    critical_mask = df['color'].isin(['red', 'yellow'])
    critical_pikets = df[critical_mask]['piket'].tolist()
    if critical_pikets:
        p = doc.add_paragraph('Следующие пикеты требуют дополнительной проверки (красные или жёлтые):')
        p.add_run(f' {", ".join(map(str, critical_pikets))}').bold = True
        doc.add_paragraph('Обратите внимание: для этих пикетов рекомендуется провести детальный анализ '
                          'и при необходимости скорректировать проект.')
    else:
        doc.add_paragraph('Все пикеты находятся в зелёной зоне. Дополнительная проверка не требуется.')

    doc.add_paragraph()
    doc.add_heading('Нормативная база', level=2)
    doc.add_paragraph('• СП 122.13330.2012 "Тоннели железнодорожные и автодорожные"')
    doc.add_paragraph('• ГОСТ 33153-2014 "Дороги автомобильные общего пользования. Тоннели автомобильные. Правила проектирования"')
    doc.add_paragraph('• ГОСТ 32836-2014 "Инженерные изыскания для строительства. Общие правила" – учтены требования к исходным данным.')

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def generate_xlsx_bytes(df):
    wb = Workbook()
    ws = wb.active
    ws.title = "Расчётные данные"
    headers = list(df.columns)
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(name='Times New Roman', size=12, bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    for row_idx, row in enumerate(df.itertuples(index=False), 2):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = Font(name='Times New Roman', size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center')
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        adjusted_width = max_length + 2
        ws.column_dimensions[col_letter].width = adjusted_width
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = thin_border
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ------------------------------------------------------------
# 4. STREAMLIT ИНТЕРФЕЙС
# ------------------------------------------------------------
st.set_page_config(page_title="Геомеханический скрининг тоннелей", layout="wide")
st.title("🚇 Геомеханический скрининг тоннелей")
st.markdown("### 1D скрининг по СП 122.13330.2012, ГОСТ 33153-2014, ГОСТ 32836-2014")

# Дисклеймер
st.warning("⚠️ Финальное проектное решение принимает главный инженер проекта (ГИП). "
           "Представленные расчёты носят рекомендательный характер и не отменяют экспертизу.")

# Инициализация состояния
if 'raw_data' not in st.session_state:
    st.session_state['raw_data'] = None
if 'results' not in st.session_state:
    st.session_state['results'] = None
if 'project_name' not in st.session_state:
    st.session_state['project_name'] = "Тоннель №1"
# Настройки модели
if 'excess_threshold' not in st.session_state:
    st.session_state['excess_threshold'] = 0.9
if 'stability_threshold' not in st.session_state:
    st.session_state['stability_threshold'] = 0.8
if 'include_plastic' not in st.session_state:
    st.session_state['include_plastic'] = False

# Боковая панель
with st.sidebar:
    st.header("Загрузка данных")
    uploaded_file = st.file_uploader("Выберите CSV или Excel файл", type=["csv", "xlsx"])
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            required = ['piket', 'H', 'B', 'c', 'phi', 'gamma', 'f', 'UGW']
            if all(col in df.columns for col in required):
                st.session_state['raw_data'] = df
                st.success(f"Загружено {len(df)} пикетов")
            else:
                st.error(f"Файл должен содержать колонки: {', '.join(required)}")
        except Exception as e:
            st.error(f"Ошибка загрузки: {e}")

    st.divider()
    st.header("Тестовые данные")
    if st.button("Загрузить тестовые данные"):
        data = {
            'piket': np.arange(0, 1000, 100),
            'H': np.random.uniform(20, 50, 10),
            'B': np.full(10, 10.0),
            'c': np.random.uniform(10, 30, 10),
            'phi': np.random.uniform(20, 35, 10),
            'gamma': np.full(10, 2.0),
            'f': np.random.uniform(2, 6, 10),
            'UGW': np.random.uniform(10, 40, 10),
        }
        st.session_state['raw_data'] = pd.DataFrame(data)
        st.success("Тестовые данные загружены")

    st.divider()
    st.header("Имя проекта")
    st.session_state['project_name'] = st.text_input("Название", st.session_state['project_name'])

    st.divider()
    st.header("Настройки модели")
    # Слайдер для порога превышения порового давления
    excess_threshold = st.slider(
        "Порог превышения порового давления (u / σ_v)",
        min_value=0.7,
        max_value=1.0,
        value=st.session_state['excess_threshold'],
        step=0.05,
        help="Чем выше значение, тем меньше пикетов будет считаться красными из-за порового давления."
    )
    st.session_state['excess_threshold'] = excess_threshold

    # Слайдер для порога устойчивости
    stability_threshold = st.slider(
        "Порог устойчивости (коэффициент запаса)",
        min_value=0.5,
        max_value=1.0,
        value=st.session_state['stability_threshold'],
        step=0.05,
        help="Чем ниже значение, тем меньше пикетов будет считаться красными из-за низкой устойчивости."
    )
    st.session_state['stability_threshold'] = stability_threshold

    # Чекбокс "Учитывать пластику в критериях красного"
    include_plastic = st.checkbox(
        "Учитывать пластику в критериях красного",
        value=st.session_state['include_plastic'],
        help="Если включено, то красный цвет присваивается при (exceeds_u и low_stability) ИЛИ (plastic и exceeds_u)."
    )
    st.session_state['include_plastic'] = include_plastic

    if st.session_state['raw_data'] is not None and st.button("▶ Выполнить расчёт", type="primary"):
        with st.spinner("Выполняется расчёт..."):
            df_full = compute_full_table(
                st.session_state['raw_data'],
                excess_threshold=st.session_state['excess_threshold'],
                stability_threshold=st.session_state['stability_threshold'],
                include_plastic=st.session_state['include_plastic']
            )
            st.session_state['results'] = df_full
        st.success("Расчёт завершён!")

# Основная область
if st.session_state['raw_data'] is None:
    st.info("Загрузите данные через боковую панель или используйте тестовые.")
else:
    st.subheader("Исходные данные")
    st.dataframe(st.session_state['raw_data'].head())

if st.session_state['results'] is not None:
    df_res = st.session_state['results']
    st.success(f"Расчёт выполнен для {len(df_res)} пикетов")

    # Вкладки для визуализации
    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["Продольный разрез", "Эпюры пикета", "Кривые CCM", "Корреляционная матрица", "Таблица аномалий"]
    )

    with tab1:
        st.plotly_chart(plot_longitudinal_profile(df_res), use_container_width=True)

    with tab2:
        piket_input = st.number_input(
            "Введите номер пикета",
            min_value=int(df_res['piket'].min()),
            max_value=int(df_res['piket'].max()),
            step=100
        )
        st.plotly_chart(plot_picket_epures(df_res, piket_input), use_container_width=True)

    with tab3:
        selected = st.multiselect(
            "Выберите пикеты для кривых CCM",
            df_res['piket'].tolist(),
            default=df_res['piket'].iloc[:3].tolist()
        )
        st.plotly_chart(plot_ccm_curves(df_res, selected), use_container_width=True)

    with tab4:
        st.plotly_chart(plot_correlation_matrix(df_res), use_container_width=True)

    with tab5:
        st.plotly_chart(plot_anomaly_table(df_res), use_container_width=True)

    # Отчёты
    st.divider()
    st.subheader("Генерация отчётов")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 Скачать DOCX отчёт"):
            docx_bytes = generate_docx_bytes(df_res, st.session_state['project_name'])
            st.download_button("Скачать DOCX", docx_bytes, file_name="report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        if st.button("📊 Скачать XLSX данные"):
            xlsx_bytes = generate_xlsx_bytes(df_res)
            st.download_button("Скачать XLSX", xlsx_bytes, file_name="report_data.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.caption("Нормативная база: СП 122.13330.2012, ГОСТ 33153-2014, ГОСТ 32836-2014")

else:
    st.info("Выполните расчёт, чтобы увидеть результаты.")
   

   
